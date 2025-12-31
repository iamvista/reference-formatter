"""
學術文獻格式整理工具 - Flask 主應用
Academic Reference Formatter - Main Flask Application
"""

from flask import Flask, render_template, request, jsonify, send_file
import os
import io
from modules.parser import ReferenceParser
from modules.api_client import APIClient
from modules.formatter import ReferenceFormatter
from config import config

# 創建應用實例
app = Flask(__name__)

# 載入配置
env = os.environ.get('FLASK_ENV', 'development')
app.config.from_object(config[env])

# 初始化模組
parser = ReferenceParser()
api_client = APIClient()

@app.route('/')
def index():
    """主頁面"""
    return render_template('index.html')

@app.route('/health')
def health_check():
    """健康檢查端點（用於監控）"""
    return jsonify({
        'status': 'healthy',
        'service': 'reference-formatter',
        'version': '1.0.0'
    }), 200

@app.route('/parse', methods=['POST'])
def parse_references():
    """
    解析文獻端點
    接收純文字文獻，返回解析和格式化結果
    """
    try:
        data = request.get_json()
        raw_text = data.get('text', '')
        format_style = data.get('format', 'apa')  # 默認使用 APA 格式
        enrich = data.get('enrich', True)  # 是否使用 API 補完資料

        # 將文獻文字按行分割（每行一條文獻）
        lines = [line.strip() for line in raw_text.split('\n') if line.strip()]

        # 解析和格式化文獻
        results = []
        for i, line in enumerate(lines):
            # 1. 解析文獻
            parsed_data = parser.parse_reference(line)

            # 2. 如果需要，使用 API 補完資料
            if enrich and (parsed_data.get('doi') or parsed_data.get('title')):
                try:
                    enriched_data = api_client.enrich_reference(parsed_data)
                    if enriched_data.get('enriched'):
                        parsed_data = enriched_data
                        status = 'enriched'
                    else:
                        status = 'complete'
                except Exception as e:
                    print(f"API 查詢失敗: {e}")
                    status = 'complete'
            else:
                status = 'complete'

            # 3. 格式化為所有格式（用於前端切換）
            formatted_refs = {}
            for style in ReferenceFormatter.get_available_styles():
                try:
                    formatted_refs[style] = ReferenceFormatter.format(parsed_data, style)
                except Exception as e:
                    formatted_refs[style] = f"[格式化失敗: {str(e)}]"

            # 4. 組合結果
            results.append({
                'id': i,
                'original': line,
                'status': status,
                'formatted': formatted_refs.get(format_style, formatted_refs['apa']),
                'formatted_all': formatted_refs,  # 所有格式的版本
                'data': parsed_data,
                'completeness': parsed_data.get('completeness', 0.0),
                'confidence': parsed_data.get('confidence', 0.0)
            })

        return jsonify({
            'success': True,
            'count': len(results),
            'references': results,
            'available_formats': ReferenceFormatter.get_available_styles()
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 400

@app.route('/export/<format_type>', methods=['POST'])
def export_references(format_type):
    """
    匯出文獻端點
    支援格式：txt, docx, bibtex, html
    """
    try:
        data = request.get_json()
        references = data.get('references', [])
        citation_style = data.get('style', 'apa')

        if format_type == 'txt':
            content = '\n\n'.join([ref.get('formatted', '') for ref in references])
            buffer = io.BytesIO(content.encode('utf-8'))
            buffer.seek(0)
            return send_file(
                buffer,
                mimetype='text/plain',
                as_attachment=True,
                download_name=f'references_{citation_style}.txt'
            )

        elif format_type == 'docx':
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            doc.add_heading('參考文獻 / References', 0)

            # 添加格式資訊
            style_names = {
                'apa': 'APA 7th Edition',
                'mla': 'MLA 9th Edition',
                'chicago': 'Chicago Manual of Style 17th Edition',
                'harvard': 'Harvard Referencing Style'
            }
            style_para = doc.add_paragraph(f'格式：{style_names.get(citation_style, citation_style.upper())}')
            style_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # 空行

            # 添加每條文獻
            for ref in references:
                p = doc.add_paragraph(ref.get('formatted', ''))
                p.paragraph_format.left_indent = Pt(36)  # 懸掛縮排
                p.paragraph_format.first_line_indent = Pt(-36)

            # 保存到記憶體
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return send_file(
                buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'references_{citation_style}.docx'
            )

        elif format_type == 'bibtex':
            # 生成 BibTeX 格式
            bibtex_entries = []
            for i, ref in enumerate(references):
                ref_data = ref.get('data', {})
                entry_type = 'article' if ref_data.get('type') == 'article' else 'misc'
                cite_key = f"ref{i+1}"

                # 構建 BibTeX 條目
                entry_lines = [f"@{entry_type}{{{cite_key},"]

                if ref_data.get('title'):
                    entry_lines.append(f"  title = {{{ref_data['title']}}},")
                if ref_data.get('authors'):
                    authors_str = ' and '.join([
                        f"{a.get('last', '')}, {a.get('first', '')}"
                        for a in ref_data['authors']
                    ])
                    entry_lines.append(f"  author = {{{authors_str}}},")
                if ref_data.get('year'):
                    entry_lines.append(f"  year = {{{ref_data['year']}}},")
                if ref_data.get('journal'):
                    entry_lines.append(f"  journal = {{{ref_data['journal']}}},")
                if ref_data.get('volume'):
                    entry_lines.append(f"  volume = {{{ref_data['volume']}}},")
                if ref_data.get('issue'):
                    entry_lines.append(f"  number = {{{ref_data['issue']}}},")
                if ref_data.get('pages'):
                    entry_lines.append(f"  pages = {{{ref_data['pages']}}},")
                if ref_data.get('doi'):
                    entry_lines.append(f"  doi = {{{ref_data['doi']}}},")
                if ref_data.get('url'):
                    entry_lines.append(f"  url = {{{ref_data['url']}}},")

                entry_lines.append("}")
                bibtex_entries.append('\n'.join(entry_lines))

            content = '\n\n'.join(bibtex_entries)
            buffer = io.BytesIO(content.encode('utf-8'))
            buffer.seek(0)

            return send_file(
                buffer,
                mimetype='text/plain',
                as_attachment=True,
                download_name='references.bib'
            )

        elif format_type == 'html':
            # 生成 HTML 格式
            style_names = {
                'apa': 'APA 7th Edition',
                'mla': 'MLA 9th Edition',
                'chicago': 'Chicago Manual of Style 17th Edition',
                'harvard': 'Harvard Referencing Style'
            }

            html_content = f"""
<!DOCTYPE html>
<html lang="zh-TW">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>參考文獻 - {style_names.get(citation_style, citation_style.upper())}</title>
    <style>
        body {{
            font-family: "Times New Roman", Times, serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.6;
        }}
        h1 {{
            text-align: center;
            font-size: 24px;
            margin-bottom: 10px;
        }}
        .style-info {{
            text-align: center;
            color: #666;
            margin-bottom: 30px;
        }}
        .reference {{
            margin-bottom: 1em;
            padding-left: 2em;
            text-indent: -2em;
        }}
        @media print {{
            body {{
                margin: 0;
                padding: 1in;
            }}
        }}
    </style>
</head>
<body>
    <h1>參考文獻 / References</h1>
    <p class="style-info">格式：{style_names.get(citation_style, citation_style.upper())}</p>
"""
            for ref in references:
                formatted = ref.get('formatted', '')
                # 處理斜體標記（*text* -> <em>text</em>）
                formatted = formatted.replace('*', '<em>', 1).replace('*', '</em>', 1)
                while '*' in formatted:
                    formatted = formatted.replace('*', '<em>', 1)
                    if '*' in formatted:
                        formatted = formatted.replace('*', '</em>', 1)

                html_content += f'    <div class="reference">{formatted}</div>\n'

            html_content += """
</body>
</html>
"""
            return html_content

        else:
            return jsonify({'error': '不支援的格式'}), 400

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 400

if __name__ == '__main__':
    # 開發模式
    app.run(debug=True, host='0.0.0.0', port=8080)
